"""
Take a list of names and make a docx with custom invites for each one.
"""

import docx

with open('../data/guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

    for name in names:
        name = name.strip()

        document.add_paragraph('It would be a pleasure to have the company of',
                               style='Heading 1')
        document.add_paragraph(name, style='Heading 1')
        document.add_paragraph('at 11010 Memory Lane on the Evening of',
                               style='Body Text')
        document.add_paragraph('April 1st', style='Body Text')
        document.add_paragraph("at 7 o'clock", style='Normal')

        document.add_page_break()

    document.save('../data/invites.docx')

    print("File has been created and saved as 'invites.docx'")
